import chevron
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

LoadFilename = "test_doc.docx"
SaveFilename = "test_doc_edit.docx"


# 参照元: https://stackoverflow.com/a/49615968
def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


if __name__ == "__main__":
    doc = Document(LoadFilename)

    # 段落を置換
    for paragraph in doc.paragraphs:
        paragraph.text = chevron.render(
            paragraph.text, {"title": "ポラーノの広場", "end": "おしまい"}
        )
        print(paragraph.text)

    # テーブル操作
    for i, table in enumerate(doc.tables):
        print("table {}".format(i))
        for row in table.rows:
            for cell in row.cells:
                print(
                    "{} -> {}".format(
                        cell.text, " ".join([hex(ord(c)) for c in cell.text])
                    )
                )

    # 参照: http://officeopenxml.com/WPtableBorders.php
    table_border = {
        "top": {
            "sz": 4,
            "color": "auto",
            "val": "dashed",
        },
        "bottom": {"sz": 4, "color": "auto", "val": "dashed"},
        "start": {"sz": 4, "color": "auto", "val": "nil"},
        "end": {"sz": 4, "color": "auto", "val": "nil"},
    }

    target_table = doc.tables[1]
    # 2行（ヘッダー含む）→17行
    for i in range(2, 17):
        row = target_table.add_row()

        cell = row.cells[0]
        cell.text = i

        cell = row.cells[1]
        cell.text = "てすとてすと"

        cell = row.cells[2]
        cell.text = "テストテスト"

        cell = row.cells[3]

        for cell in row.cells:
            set_cell_border(cell, **table_border)

        # print(dir(docx.text.run.WD_BREAK))
        # cell.text = "{}行目\n改行テスト".format(i)
        # cell.run.add_break(docx.text.run.WD_BREAK.LINE)

    # テーブル操作
    for i, table in enumerate(doc.tables):
        print("table {}".format(i))
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)

                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                print(tcBorders)
                if tcBorders is None:
                    tcBorders = OxmlElement("w:tcBorders")
                    tcPr.append(tcBorders)

                # for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
                #     edge_data = kwargs.get(edge)

    doc.save(SaveFilename)
